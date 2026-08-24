from pathlib import Path
import re
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from bidi.algorithm import get_display
import arabic_reshaper

ROOT = Path(__file__).resolve().parent
SRC = ROOT / 'source' / 'chapter-01.md'
DOCX = ROOT / 'docx' / 'فصل-01-تحول-دیجیتال-2-0.docx'
PDF = ROOT / 'pdf' / 'فصل-01-تحول-دیجیتال-2-0.pdf'
FONT = '/usr/share/fonts/truetype/noto/NotoNaskhArabic-Regular.ttf'
FONT_BOLD = '/usr/share/fonts/truetype/noto/NotoNaskhArabic-Bold.ttf'

for p in (DOCX.parent, PDF.parent): p.mkdir(parents=True, exist_ok=True)


def clean(s):
    s = re.sub(r'cite[^]*', '', s)
    s = re.sub(r'\*\*(.*?)\*\*', r'\1', s)
    s = s.replace('—', ' — ')
    return s.strip()


def parse_md(text):
    lines = text.replace('\r\n','\n').split('\n')
    blocks=[]; i=0
    while i < len(lines):
        line=lines[i].strip()
        if not line: i+=1; continue
        if line.startswith('|'):
            rows=[]
            while i<len(lines) and lines[i].strip().startswith('|'):
                cells=[clean(x) for x in lines[i].strip().strip('|').split('|')]
                if not all(set(c.replace('-','').replace(':','').strip())==set() for c in cells): rows.append(cells)
                i+=1
            if len(rows)>1: blocks.append(('table',rows))
            continue
        m=re.match(r'^(#{1,3})\s+(.*)$', line)
        if m: blocks.append(('h',len(m.group(1)),clean(m.group(2)))); i+=1; continue
        if line.startswith('>'):
            q=[]
            while i<len(lines) and lines[i].strip().startswith('>'):
                q.append(clean(lines[i].strip()[1:].strip())); i+=1
            blocks.append(('quote',' '.join(q))); continue
        if line.startswith('**') and line.endswith('**'):
            blocks.append(('callout',clean(line))); i+=1; continue
        if re.match(r'^\d+\.\s',line) or line.startswith('- '):
            items=[]
            while i<len(lines) and (re.match(r'^\d+\.\s',lines[i].strip()) or lines[i].strip().startswith('- ')):
                items.append(clean(re.sub(r'^(?:\d+\.|-)\s*','',lines[i].strip()))); i+=1
            blocks.append(('list',items)); continue
        para=[line]; i+=1
        while i<len(lines) and lines[i].strip() and not lines[i].strip().startswith(('#','|','>','- ')) and not re.match(r'^\d+\.\s',lines[i].strip()):
            para.append(lines[i].strip()); i+=1
        blocks.append(('p',clean(' '.join(para))))
    return blocks


def set_rtl(paragraph):
    pPr=paragraph._p.get_or_add_pPr(); bidi=OxmlElement('w:bidi'); bidi.set(qn('w:val'),'1'); pPr.append(bidi)
    paragraph.alignment=WD_ALIGN_PARAGRAPH.RIGHT


def add_run(paragraph,text,bold=False,size=12):
    r=paragraph.add_run(text); r.bold=bold; r.font.name='Noto Naskh Arabic'; r.font.size=Pt(size)
    r._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Naskh Arabic')
    return r


def build_docx(blocks):
    d=Document(); sec=d.sections[0]; sec.top_margin= Cm(2.4); sec.bottom_margin=Cm(2.3); sec.right_margin=Cm(2.5); sec.left_margin=Cm(2.5)
    styles=d.styles['Normal']; styles.font.name='Noto Naskh Arabic'; styles.font.size=Pt(12); styles.paragraph_format.line_spacing=1.35; styles.paragraph_format.space_after=Pt(7)
    for b in blocks:
        if b[0]=='h':
            p=d.add_paragraph(); set_rtl(p); add_run(p,b[2],True,18 if b[1]==1 else (15 if b[1]==2 else 13)); p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(6)
        elif b[0]=='quote' or b[0]=='callout':
            p=d.add_paragraph(); set_rtl(p); add_run(p,b[1],True,12); p.paragraph_format.left_indent=Cm(.5); p.paragraph_format.right_indent=Cm(.5)
        elif b[0]=='p':
            p=d.add_paragraph(); set_rtl(p); add_run(p,b[1],False,12)
        elif b[0]=='list':
            for idx,item in enumerate(b[1],1):
                p=d.add_paragraph(); set_rtl(p); add_run(p,f'{idx}. {item}',False,12)
        elif b[0]=='table':
            rows=b[1]; t=d.add_table(rows=len(rows),cols=len(rows[0])); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
            for ri,row in enumerate(rows):
                for ci,val in enumerate(row):
                    cell=t.cell(ri,ci); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; cell.text=''
                    p=cell.paragraphs[0]; set_rtl(p); add_run(p,val,ri==0,10.5)
            d.add_paragraph()
    d.save(DOCX)


def rtl(s): return get_display(arabic_reshaper.reshape(s))

def build_pdf(blocks):
    pdfmetrics.registerFont(TTFont('NotoNaskh',FONT)); pdfmetrics.registerFont(TTFont('NotoNaskhBold',FONT_BOLD if Path(FONT_BOLD).exists() else FONT))
    doc=SimpleDocTemplate(str(PDF),pagesize=A4,rightMargin=2.5*cm,leftMargin=2.5*cm,topMargin=2.4*cm,bottomMargin=2.3*cm,title='فصل اول — مبانی تحول دیجیتال',author='AI Future Tech Radar')
    normal=ParagraphStyle('n',fontName='NotoNaskh',fontSize=11.5,leading=18,alignment=2,spaceAfter=7)
    h1=ParagraphStyle('h1',parent=normal,fontName='NotoNaskhBold',fontSize=17,leading=25,spaceBefore=12,spaceAfter=7)
    h2=ParagraphStyle('h2',parent=normal,fontName='NotoNaskhBold',fontSize=14.5,leading=22,spaceBefore=10,spaceAfter=6)
    h3=ParagraphStyle('h3',parent=normal,fontName='NotoNaskhBold',fontSize=12.5,leading=20,spaceBefore=8,spaceAfter=5)
    story=[]
    for b in blocks:
        if b[0]=='h': story.append(Paragraph(rtl(b[2]), h1 if b[1]==1 else (h2 if b[1]==2 else h3)))
        elif b[0] in ('quote','callout'): story += [Paragraph(rtl(b[1]),normal),Spacer(1,4)]
        elif b[0]=='p': story.append(Paragraph(rtl(b[1]),normal))
        elif b[0]=='list':
            for idx,item in enumerate(b[1],1): story.append(Paragraph(rtl(f'{idx}. {item}'),normal))
        elif b[0]=='table':
            data=[[Paragraph(rtl(x),normal) for x in row] for row in b[1]]
            t=Table(data,repeatRows=1,hAlign='RIGHT'); t.setStyle(TableStyle([('FONTNAME',(0,0),(-1,-1),'NotoNaskh'),('BACKGROUND',(0,0),(-1,0),None),('GRID',(0,0),(-1,-1),.4,None),('VALIGN',(0,0),(-1,-1),'MIDDLE'),('ALIGN',(0,0),(-1,-1),'RIGHT'),('RIGHTPADDING',(0,0),(-1,-1),6),('LEFTPADDING',(0,0),(-1,-1),6)])); story += [t,Spacer(1,8)]
    doc.build(story)

if __name__=='__main__':
    blocks=parse_md(SRC.read_text(encoding='utf-8')); build_docx(blocks); build_pdf(blocks); print(f'generated {DOCX} and {PDF}')
